"""
utils.py - Module for Utility Functions.

This module provides utility functions used in the PowerPoint presentation creation process,
including formatting messages with colors, setting up a debug logger, and reading Word files.

Module Functions:
    - format_prompt(message):
        Format the given message as a cyan-colored prompt.

    - format_info(message):
        Format the given message as a green-colored information.

    - format_warning(message):
        Format the given message as a yellow-colored warning.

    - format_error(message):
        Format the given message as a red-colored error.

    - setup_debug_logger(debug_mode):
        Setup and configure a debug logger.

    - read_word_file(file_path, logger):
        Read the content of a Word file.

    - read_big_word_file(file_path, chunk_size, logger):
        Read a big Word file in chunks.

    - ensure_list(result):
        Ensure the result is a list.

Dependencies:
    - logging
    - docx (from python-docx package)
    - colorama

"""
import re
import logging
from docx import Document
from colorama import Fore, Style


def format_prompt(message):
    """Format the given message as a cyan-colored prompt.

    Args:
        message (str): The message to be formatted.

    Returns:
        str: The formatted message.
    """
    return Fore.CYAN + message +"> " + Style.RESET_ALL

def format_info(message):
    """Format the given message as a green-colored information.

    Args:
        message (str): The message to be formatted.

    Returns:
        str: The formatted message.
    """
    return Fore.GREEN + message + Style.RESET_ALL

def format_warning(message):
    """Format the given message as a yellow-colored warning.

    Args:
        message (str): The message to be formatted.

    Returns:
        str: The formatted message.
    """
    return Fore.YELLOW + "Warning: " + message + Style.RESET_ALL

def format_error(message):
    """Format the given message as a red-colored error.

    Args:
        message (str): The message to be formatted.

    Returns:
        str: The formatted message.
    """
    return Fore.RED + "Error: " + message + Style.RESET_ALL

def setup_debug_logger(debug_mode):
    """Setup and configure a debug logger.

    Args:
        debug_mode (bool): Whether to enable console logging.

    Returns:
        logging.Logger: The configured logger object.
    """
    # Configure the logger
    logger = logging.getLogger("debug_logger")
    logger.setLevel(logging.DEBUG)

    # Create a file handler that logs to a file
    file_handler = logging.FileHandler("debug_log.txt", encoding='utf-8')
    file_handler.setLevel(logging.DEBUG)

    # Create a console handler that prints debug messages to the console
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.DEBUG)

    # Create a formatter for the log messages
    formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')

    # Attach the formatter to the handlers
    file_handler.setFormatter(formatter)
    console_handler.setFormatter(formatter)

    # Add the handlers to the logger
    logger.addHandler(file_handler)
    # Add the console handler only if in debug mode
    if debug_mode:
        logger.addHandler(console_handler)

    return logger

def read_word_file(file_path,logger):
    """Read the content of a Word file.

    Args:
        file_path (str): The path of the Word file to read.
        logger (logging.Logger): The logger object to record any errors.

    Returns:
        str or None: The full text content of the Word file, or None if an error occurs.
    """
    try:
        doc = Document(file_path)
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        return full_text
    except Exception as e:
        print(format_error("Error reading the Word file: ")+ f"{e}")
        logger.error("Error reading the Word file:", e)
        return None

def read_big_word_file(file_path,chunk_size,logger):
    """Read a big Word file in chunks.

    Args:
        file_path (str): The path of the Word file to read.
        chunk_size (int): The maximum size of each chunk (in words).
        logger (logging.Logger): The logger object to record any errors.

    Returns:
        list or None: A list of chunks containing the full text content of the Word file,
        or None if an error occurs.
    """
    try:
        doc = Document(file_path)
        full_texts = []
        current_text = ""

        for paragraph in doc.paragraphs:
            current_paragraph_text = paragraph.text + "\n"

            # Check if adding the current paragraph will exceed 1000 words
            if len(current_text.split()) + len(current_paragraph_text.split()) > chunk_size:
                # If so, add the current_text to the list and start a new element
                full_texts.append(current_text)
                current_text = current_paragraph_text
            else:
                # Otherwise, continue adding the paragraph to the current_text
                current_text += current_paragraph_text

        # Add any remaining text to the list
        if current_text:
            full_texts.append(current_text)

        return full_texts
    except Exception as e:
        print(format_error("Error reading the Big Word file: ")+ f"{e}")
        logger.error("Error reading the Big Word file:", e)
        return None


def ensure_list(result):
    """Ensure the result is a list.

    Args:
        result (dict or list): The result to be converted to a list.

    Returns:
        list: The result as a list.

    Raises:
        TypeError: If the result is neither a dictionary nor a list.
    """
    if isinstance(result, dict):
        return process_dict_list([result])
    elif isinstance(result, list):
        return process_dict_list(result)
    else:
        raise TypeError("Result must be a dictionary or a list.")

def extract_json_from_string(input_string):
    # Define a pattern for finding JSON content
    json_pattern = re.compile(r'\[.*\]', re.DOTALL)
    # Search for the pattern in the input string
    match = json_pattern.search(input_string)

    # Check if a match is found
    if match:
        # Extract and return the matched JSON content
        json_content = match.group(0)
        return json_content
    else:
        # Define a pattern for finding JSON content
        json_pattern = re.compile(r'\{.*\}', re.DOTALL)
        # Search for the pattern in the input string
        match = json_pattern.search(input_string)
        if match:
            # Extract and return the matched JSON content
            json_content = match.group(0)
            return json_content
        else:
            return None

def process_content(input_data):
    if isinstance(input_data, str):
        return input_data
    elif isinstance(input_data, dict):
        return '\n'.join([f"{key}: {value}" for key, value in input_data.items()])
    elif isinstance(input_data, list):
        output = []
        for item in input_data:
            if isinstance(item, str):
                output.append(item)
            elif isinstance(item, dict):
                output.extend([f"{key}: {value}" for key, value in item.items()])
            else:
                return "{fill the placeholder}"
        return '\n'.join(output)
    else:
        return "{fill the placeholder}"

def process_dict_list(dict_list):
    for dictionary in dict_list:
        if "content" in dictionary:
            dictionary["content"] = process_content(dictionary["content"])
    return dict_list